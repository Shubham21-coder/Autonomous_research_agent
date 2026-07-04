import os
from dotenv import load_dotenv
import json
from fastapi.responses import StreamingResponse, Response
import io

# CRITICAL: This must be executed at the absolute top of the file 
# before any other modules or engine dependencies are imported/instantiated.
load_dotenv()

from fastapi import FastAPI, Depends, HTTPException, BackgroundTasks, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from sqlalchemy.orm import Session
from pydantic import BaseModel
from database import SessionLocal, User, SearchHistory, get_password_hash, verify_password, create_access_token
from agent import AutonomousAgentEngine
import jwt
import markdown
import pdfkit
from docx import Document

app = FastAPI(title="Autonomous AI Agent API")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_credentials=True, allow_methods=["*"], allow_headers=["*"])

oauth2_scheme = OAuth2PasswordBearer(tokenUrl="api/auth/login")

# Now this can safely pull HF_TOKEN from the fully initialized environment
agent_engine = AutonomousAgentEngine()

# --- Dependencies ---
def get_db():
    db = SessionLocal()
    try: yield db
    finally: db.close()

def get_current_user(token: str = Depends(oauth2_scheme), db: Session = Depends(get_db)):
    try:
        payload = jwt.decode(token, os.getenv("JWT_SECRET_KEY", "super_secret_agent_key_change_in_production"), algorithms=["HS256"])
        user = db.query(User).filter(User.username == payload.get("sub")).first()
        if user is None: raise HTTPException(status_code=401)
        return user
    except: raise HTTPException(status_code=401, detail="Invalid credentials")

# --- Schemas ---
class UserCreate(BaseModel):
    username: str
    password: str

class QueryRequest(BaseModel):
    query: str

class ExportRequest(BaseModel):
    markdown_text: str
    format: str

# --- Endpoints ---
@app.post("/api/auth/signup")
def signup(user: UserCreate, db: Session = Depends(get_db)):
    if db.query(User).filter(User.username == user.username).first():
        raise HTTPException(status_code=400, detail="Username registered")
    new_user = User(username=user.username, hashed_password=get_password_hash(user.password))
    db.add(new_user)
    db.commit()
    return {"message": "User created. You may now login."}

@app.post("/api/auth/login")
def login(form_data: OAuth2PasswordRequestForm = Depends(), db: Session = Depends(get_db)):
    user = db.query(User).filter(User.username == form_data.username).first()
    if not user or not verify_password(form_data.password, user.hashed_password):
        raise HTTPException(status_code=400, detail="Incorrect username or password")
    token = create_access_token(data={"sub": user.username})
    return {"access_token": token, "token_type": "bearer"}


@app.post("/api/research/run")
async def execute_research(request: QueryRequest, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    async def event_generator():
        # Yield Stage 1
        yield f"data: {json.dumps({'stage': 'Plan', 'status': 'Active'})}\n\n"
        plan = agent_engine.plan(current_user.id, request.query)
        
        # Yield Stage 2
        yield f"data: {json.dumps({'stage': 'Deep Search', 'status': 'Active'})}\n\n"
        raw_data = await agent_engine.execute_searches(plan)
        enriched_data = await agent_engine.deeply_scrape_results(raw_data)
        
        # Yield Stage 3
        yield f"data: {json.dumps({'stage': 'Rank & Embed', 'status': 'Active'})}\n\n"
        context = agent_engine.rank_and_extract(current_user.id, request.query, enriched_data)
        
        # Yield Stage 4
        yield f"data: {json.dumps({'stage': 'Synthesis', 'status': 'Active'})}\n\n"
        final_report = agent_engine.synthesize_and_verify(request.query, context)
        
        db.add(SearchHistory(user_id=current_user.id, query=request.query, result_path="Markdown Data Saved"))
        db.commit()
        
        # Yield Final Output
        yield f"data: {json.dumps({'report': final_report})}\n\n"

    return StreamingResponse(event_generator(), media_type="text/event-stream")

import io
import markdown
import pdfkit
from docx import Document
from fastapi import Response
from pydantic import BaseModel

class ExportRequest(BaseModel):
    markdown_text: str
    format: str

@app.post("/api/research/export")
async def export_report(request: ExportRequest, current_user: User = Depends(get_current_user)):
    if request.format == "pdf":
        html = markdown.markdown(request.markdown_text)
        try:
            # Requires wkhtmltopdf installed on the host machine
            pdf_bytes = pdfkit.from_string(html, False)
            return Response(
                content=pdf_bytes, 
                media_type="application/pdf", 
                headers={"Content-Disposition": "attachment; filename=ARIA_Report.pdf"}
            )
        except Exception as e:
            return Response(status_code=500, content=f"PDF Engine Error. Ensure wkhtmltopdf is installed. Log: {str(e)}")
    
    elif request.format == "docx":
        doc = Document()
        for line in request.markdown_text.split('\n'):
            if line.startswith('# '):
                doc.add_heading(line.replace('# ', ''), level=1)
            elif line.startswith('## '):
                doc.add_heading(line.replace('## ', ''), level=2)
            elif line.startswith('### '):
                doc.add_heading(line.replace('### ', ''), level=3)
            elif line.strip() == '':
                continue
            else:
                doc.add_paragraph(line)
        
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return Response(
            content=file_stream.read(), 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
            headers={"Content-Disposition": "attachment; filename=ARIA_Report.docx"}
        )